"""
Generate the RPPR report for ASAP AViDD U19 year 1.

TODO:
* Switch to using markdown2docx in future, which enables better control over mappings from Markdown styles to Word Styles?
  https://pypi.org/project/Markdown2docx/

"""

#
# Date range
#

import datetime
#reporting_period_start = datetime.date.fromisoformat('2022-05-01')
#reporting_period_end = datetime.date.fromisoformat('2023-04-30')
reporting_period_start = datetime.date.fromisoformat('2023-05-01')
reporting_period_end = datetime.date.fromisoformat('2024-04-30')
grant_id = 'NIH U19 AI171399' 

# -*- coding: utf-8 -*-

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, Cm

# Date functions
def published_during_reporting_period(paper):
    """Return True if paper was published during the grant reporting period.
    """
    try:
        if reporting_period_start <= paper['published']['dates']['published'] <= reporting_period_end:
            return True
    except Exception as e:
        pass
    return False

def accepted_during_reporting_period(paper):
    """Return True if paper was published during the grant reporting period.
    """
    try:
        if reporting_period_start <= paper['published']['dates']['accepted'] <= reporting_period_end:
            return True
    except Exception as e:
        pass
    return False

def preprinted_during_reporting_period(paper):
    """Return True if paper was published during the grant reporting period.
    """
    try:
        if reporting_period_start <= paper['preprint']['date'] <= reporting_period_end:
            return True
    except Exception as e:
        pass
    return False

#
# Define Markdown to docx renderer
#

import mistune

class MathBlockGrammar(mistune.BlockGrammar):
    import re
    block_math = re.compile(r"^\$\$(.*?)\$\$", re.DOTALL)


class MathBlockLexer(mistune.BlockLexer):
    default_rules = ['block_math'] + mistune.BlockLexer.default_rules

    def __init__(self, rules=None, **kwargs):
        if rules is None:
            rules = MathBlockGrammar()
        super(MathBlockLexer, self).__init__(rules, **kwargs)

    def parse_block_math(self, m):
        """Parse a $$math$$ block"""
        self.tokens.append({'type': 'block_math', 'text': m.group(1)})


class MarkdownWithMath(mistune.Markdown):
    def __init__(self, renderer, **kwargs):
        kwargs['block'] = MathBlockLexer
        super(MarkdownWithMath, self).__init__(renderer, **kwargs)

    def output_block_math(self):
        return self.renderer.block_math(self.token['text'])

# Generate the code to render the document
# TODO: Convert this to actually execute the code on a provided document object instead of jut returning it
class PythonDocxRenderer(mistune.Renderer):
    def __init__(self, **kwds):
        super(PythonDocxRenderer, self).__init__(**kwds)
        self.table_memory = []
        self.img_counter = 0
        self.list_level_counter = 0

    def header(self, text, level, raw):
        return "p = document.add_heading('', %d)\n" % (level) + text

    def paragraph(self, text):
        if 'add_picture' in text:
            return text
        add_break = '' if text.endswith(':")\n') else 'p.add_run().add_break()'
        add_break = '' # DEBUG
        return '\n'.join(('p = document.add_paragraph()', text, add_break)) + '\n'

    def list(self, body, ordered):
        #return body + '\np.add_run().add_break()\n'
        return body + ''
        
    def list_item(self, text):
        return '\n'.join(("p = document.add_paragraph('', style = 'BasicUserList')", text))

    def table(self, header, body):
        import itertools
        number_cols = header.count('\n') - 2
        number_rows = int(len(self.table_memory) / number_cols)
        cells = ["table.rows[%d].cells[%d].paragraphs[0]%s\n" % (i, j, self.table_memory.pop(0)[1:]) for i, j in itertools.product(range(number_rows), range(number_cols))]
        return '\n'.join(["table = document.add_table(rows=%d, cols=%d, style = 'BasicUserTable')" % (number_rows, number_cols)] + cells) + 'document.add_paragraph().add_run().add_break()\n'

    def table_cell(self, content, **flags):
        self.table_memory.append(content)
        return content

    # SPAN LEVEL
    def text(self, text):
        import re
        text = re.sub("\s+", " ", text)
        text = re.sub('"', '\\"', text)        
        return "p.add_run(\"%s\")\n" % text

    def emphasis(self, text):
        return text[:-1] + '.italic = True\n'

    def double_emphasis(self, text):
        return text[:-1] + '.bold = True\n'

    def codespan(self, text):
        return "p.add_run(\"%s\", style=\"CodeSpan\")\n" % text

    def block_code(self, code, language):
        code = code.replace('\n', '\\n')
        return "p = document.add_paragraph()\np.add_run(\"%s\")\np.style = 'BlockCode'\np.add_run().add_break()\n" % code

    def link(self, link, title, content):
        #return f'p.add_run("{content} [")\np.add_run("{link}", style="CodeSpan")\np.add_run("]")\n'
        return f'p.add_run("{content} [{link}]")\n'

    def autolink(self, link, is_email=False):
        # TODO: Render autolinks
        return f'p.add_run("{link}", style="CodeSpan")\n'

    def inline_html(self, text):
        # DEBUG
        return ""

    def image(self, src, title, alt_text):
        return '\n'.join((
            "p = document.add_paragraph()",
            "p.alignment = WD_ALIGN_PARAGRAPH.CENTER",
            "p.space_after = Pt(18)",
            "run = p.add_run()",
            "run.add_picture(\'%s\')" % src if "tmp" in src else "run.add_picture(\'%s\', width=Cm(15))" % src,
            "run.add_break()",
            "run.add_text(\'%s\')" % alt_text,
            "run.font.italic = True",
            "run.add_break()"
            )) + '\n'

    def hrule(self):
        return "document.add_page_break()\n"

    def block_math(self, text):
        import sympy
        if not os.path.exists('tmp'):
            os.makedirs('tmp')
        filename = 'tmp/tmp%d.png' % self.img_counter
        self.img_counter = self.img_counter + 1
        sympy.preview(r'$$%s$$' % text, output='png', viewer='file', filename=filename, euler=False)
        return self.image(filename, None, "Equation " + str(self.img_counter - 1))

def get_components(yaml_filename='specific-aims.yaml'):
    """
    Retrieve all component short and full names

    Returns
    -------
    components : dict of str : str
        components[short_name] is the long name of that component
    """
    # Read Specific Aims from YAML file 'specific-aims.yaml'
    import yaml
    with open(yaml_filename, 'r') as f:
        components = yaml.load(f, Loader=yaml.SafeLoader)

    return components    

def get_specific_aims(component_name):
    """
    Retrieve the text for the Specific Aims section of the Progress Report for the specified Project or Core

    -----------
    component_name: str
        Name of the project or core to generate a report for.
    """
    # Read Specific Aims from YAML file 'specific-aims.yaml'
    import yaml
    specific_aims_filename = 'specific-aims.yaml'
    with open(specific_aims_filename, 'r') as f:
        specific_aims = yaml.load(f, Loader=yaml.SafeLoader)
    # Check to make sure the Specific Aims can be found for this Project or Core
    if not component_name in specific_aims:
        raise ValueError(f'No Specific Aims found for component {component_name} in {specific_aims_filename}')

    if specific_aims[component_name]['funded_aims_modified']:
        return specific_aims[component_name]['aims']
    else:
        return "The Aims of this component have not been modified from the original, competing application."

def component_contributed_to_output(component_shortname, output):
    """
    Return True if specified component contributed to this research output
    """
    projects = output.get('projects', list())
    cores = output.get('cores', list())
    if component_shortname in (projects + cores):
        return True
    else:
        return False

def filter_events_to_reporting_period(output):
    """
    Remove events not in reporting period

    Return True if at least one event remains; else False
    """
    import datetime
    events = list()
    for event in output['events']:
        text_date = str(event['date'])
        event_date = datetime.date.fromisoformat(text_date)
        if reporting_period_start <= event_date <= reporting_period_end:
            events.append(event)
    output['events'] = events

    if len(events) > 0:
        return True
    else:
        return False

def render_products_to_markdown(yaml_filepath, component_shortname):
    """
    Render list of products to Markdown from a given YAML file or folder of YAML files.

    Parameters
    ----------
    yaml_filepath : str
        Path to a YAML file or folder of YAML files containing products to render
    component_shortname : str
        The short name of the Project/Core to render research outputs for

    Returns
    -------
    markdown_text : str
        Markdown text for the research outputs of this Project/Core

    """
    # Get all research outputs
    outputs = get_components(yaml_filename=yaml_filepath)

    markdown_text = ""

    # Flatten to list
    if hasattr(outputs, 'items'):
        outputs = list(outputs.values())

    # Iterate over outputs and include those that match this component
    for output in outputs:
        # Check if the component contributed to this output        
        if component_contributed_to_output(component_shortname, output) and filter_events_to_reporting_period(output):

            # Name
            markdown_text += f"**{output['name']}**"
            if output.get('status', None) == 'draft':
                markdown_text += ' [DRAFT]'
            markdown_text += '\n\n'

            # Description
            if 'description' in output:
                markdown_text += f"{output['description']}\n\n"

            # Render links
            if 'permalink' in output:
                markdown_text += f"{output['permalink']}\n\n"
            else:
                for link in output['links']:
                    markdown_text += f"* _{link['name']}_\n\n"
                    markdown_text += f"{link['url']}\n\n"

            # Render contributing Projects and Cores
            markdown_text += 'Contributing Projects and Cores: '
            markdown_text += ', '.join(output.get('projects', []) + output.get('cores', []))
            markdown_text += '\n\n'

            # Render events
            for event in output['events']:
                markdown_text += f"* {event['date']} : {event['description']}\n"                 

            markdown_text += '\n'

    return markdown_text

def render_research_outputs(component_shortname):
    """
    Render research outputs for this Project/Core to Markdown.

    Parameters
    ----------
    component_shortname : str
        The short name of the Project/Core to render research outputs for

    Returns
    -------
    markdown_text : str
        Markdown text for the research outputs of this Project/Core
    """

    # Resources to process
    resource_filenames = [
        'targeting_opportunities.yaml', # project 1
        'circulating_variants.yaml', # project 1
        'molecules.yaml', # projects 3, 4, 5
        'TCPs.yaml', # project 5
        'TPPs.yaml', # project 6
        'assay_cascades.yaml', # Projects 3, 4, 5
        'assay_protocols.yaml', # Biochemical Assay and Antiviral Core
    ]

    # TODO: add counts for Structural Biology, Biochemical Assay, and Antiviral Core
    # TODO: Check antiviral core progress report

    markdown_text = ""
    for resource_filename in resource_filenames:
        import os
        resource_filepath = os.path.join('../data/outputs', resource_filename)
        markdown_text += render_products_to_markdown(resource_filepath, component_shortname)

    return markdown_text

def render_TEP_outputs(component_shortname):
    """
    Render TEP outputs for this Project/Core to Markdown.

    Parameters
    ----------
    component_shortname : str
        The short name of the Project/Core to render research outputs for

    Returns
    -------
    markdown_text : str
        Markdown text for the research outputs of this Project/Core
    """

    markdown_text = ""

    # Get all TEPs
    import glob
    yaml_filenames = glob.glob('../data/outputs/TEPs/*.yaml')
    for yaml_filename in yaml_filenames:
        TEP = get_components(yaml_filename=yaml_filename)['TEP']
        # Skip TEPs without reported resources
        if not 'resources' in TEP:
            continue

        # Assess resources for inclusion
        for resource in TEP['resources']:
            if 'date' not in resource:
                print(resource)

            # Check if this resource is in the reporting period
            try:
                text_date = str(resource['date'])
                import datetime
                event_date = datetime.date.fromisoformat(text_date)
            except ValueError as e:
                print(e)
                continue
            if not (reporting_period_start <= event_date <= reporting_period_end):
                continue

            # Check if this component contributed to this output
            if not component_contributed_to_output(component_shortname, resource):
                continue

            # Don't report things without URLs
            #if not 'url 'in resource:
            #    continue

            # Render the resource
            markdown_text += f"**{resource['name']}**\n\n"
            markdown_text += f"*Date completed:* {resource['date']}\n\n"

            # Description
            if 'description' in resource:
                description = str(resource['description'])
                markdown_text += f"{description}\n\n"

            # Render links
            if 'id' in resource:
                url = 'http://asapdiscovery.org/outputs/target-enabling-packages/#' + resource['id']
                markdown_text += f"{url}\n\n"
            elif 'url' in resource:
                markdown_text += f"{resource['url']}\n\n"
            else:
                markdown_text += f"Web link pending\n\n"

            # Render contributing Projects and Cores
            markdown_text += 'Contributing Projects and Cores: '
            markdown_text += ', '.join(resource.get('projects', []) + resource.get('cores', []))
            markdown_text += '\n\n'

    #for index, line in enumerate(markdown_text.split('\n')):
    #    print(f'{index:8d}: {line}')

    return markdown_text

def generate_progress_report(component_shortname, component, output_path):
    """
    Generate a project report for the specified Project or Core

    Parameters:
    -----------
    component_shortname : str
        Component short name key used in data tables
    component : dict
        Component metadata containing fields
            name: full neame
            type: Project or Core
            significance: the significance of this Project or Core
            funded_aims_modified: if Specific Aims have been modified from submitted proposal
            aims: Specific Aims
    output_path : str
        Name of output path to write report to.

    """
    accomplishments = get_components(yaml_filename='asap-year-1-accomplishments.yaml')
    plans = get_components(yaml_filename='asap-year-1-plans.yaml')

    # Use docx to generate Word documents on continuation pages directly.
    # See https://python-docx.readthedocs.io/
    import docx

    # Load the template and extract styles
    docx_template_filename = "2590_continuation-template.docx"
    document = docx.Document(docx_template_filename)
    #styles = document.styles
    #from docx.enum.style import WD_STYLE_TYPE
    #paragraph_styles = [
    #    s for s in styles if s.type == WD_STYLE_TYPE.PARAGRAPH
    #]
    #for style in paragraph_styles:
    #    print(style.name)

    # Configure the header
    document.sections[0].header.paragraphs[0].add_run('Chodera, John Damon').bold = True

    # Construct markdown text for content of the report
    markdown_text = ""

    # Add heading for Project/Core name
    document.add_heading(component['name'], 0)

    # Add requested narrative information to go before the standard Progress Report material
    if component_shortname == 'Administrative Core':
        #
        # Administrative core
        #
        
        # A manually-asembled narrative goes in front
        pass        

    elif component['type'] == 'Project':
        #
        # Project-specific narrative
        #

        # Significant changes in the Specific Aims
        # TODO: Should this section describe any changes in approach during the past year?
        markdown_text += "# Significant changes in the Specific Aims\n\n"
        if component['funded_aims_modified']:
            markdown_text += 'The Specific Aims have been modified from the original, competing application as described in "A. Specific Aims" below.\n\n'
        else:
            markdown_text += 'The Specific Aims have not been modified from the original, competing application.\n\n'
            
        # Significance of the work
        markdown_text += '# Significance of the work\n\n'
        markdown_text += component['significance'] + '\n\n'

        # Significance of the work
        markdown_text += '# Product development milestones\n\n'
        if 'product_development_milestones' in component:
            markdown_text += component['product_development_milestones'] + '\n\n'
        else:
            markdown_text += 'This component does not have any product development milestones because it only supports the discovery stage.\n\n'

        # Significant Project-generated resources
        markdown_text += "# Significant Project-Generated Resources\n\n"
        markdown_text += render_research_outputs(component_shortname)
        markdown_text += render_TEP_outputs(component_shortname)

    elif component['type'] == 'Core':
        #
        # Core-specific narrative (except for Administrative Core, which has text prepended manually)
        #
        
        # Individual Research Projects served and activities performed and/or completed
        markdown_text += "# Individual Research Projects served and activities performed and/or completed\n\n"
        if 'projects_served' in component:
            markdown_text += component['projects_served'] + '\n\n'
        # TODO: List activities performed and/or completed

        # Significant Core-generated resources (if any)
        markdown_text += "# Significant Core-Generated Resources\n\n"
        markdown_text += render_research_outputs(component_shortname)
        markdown_text += render_TEP_outputs(component_shortname)

    # Page separator
    markdown_text += "---\n\n"

    # A. Specific Aims
    markdown_text += '# A. Specific Aims\n\n'
    if component['funded_aims_modified']:        
        markdown_text += component['aims'] + '\n\n'
    else:
        markdown_text += 'The Specific Aims have not been modified from the original, competing application.\n\n'

    markdown_text += "---\n\n"

    # B. Studies and Results
    markdown_text += '# B. Studies and Results\n\n'
    markdown_text += 'Significant accomplishments include:\n\n'
    markdown_text += accomplishments[component_shortname] + '\n\n'
    markdown_text += f'Other major results and outputs from this {component["type"]} are listed in Significant Project Generated Resources and have been posted online.\n\n'
    # Include statistics
    if component_shortname == 'Structural Biology Core':
        spreadsheets = {
            'SARS-CoV-2 Mpro protease' : 'SARS_Mpro_SBC_Analysis.xlsx',
            'MERS-CoV Mpro protease' : 'MERS_Mpro_SBC_Analysis.xlsx',
            'SARS-CoV-2 nsp3 Mac1 macrodomain' : 'Nsp3_Mac1_SBC_Analysis.xlsx',
            }

        for target, filename in spreadsheets.items():
            import pandas as pd
            sheet = pd.read_excel(f'structural-biology-core-data/{filename}', sheet_name='Experiment_Summary')
            markdown_text += f'For **{target}**, the following experiments have been conducted during the reporting period:\n\n'
            # Count the number of times each value appears in 'Experiment Status' column and collect counts into a dict
            counts = sheet['Experiment Status'].value_counts().to_dict()
            for count, name in counts.items():
                markdown_text += f'* {count} : {name}\n'
            markdown_text += '\n'

    markdown_text += "---\n\n"

    # C. Significance
    markdown_text += '# C. Significance\n\n'
    markdown_text += component['significance'] + '\n\n'
    markdown_text += "---\n\n"

    # D. Plans
    markdown_text += '# D. Plans\n\n'
    markdown_text += 'Plans for the next project period include:\n\n'
    markdown_text += plans[component_shortname] + '\n\n'
    markdown_text += "---\n\n"

    # Human Subjects
    markdown_text += "# Human Subjects\n\n"
    markdown_text += 'Not applicable\n\n'

    # Inclusion of Women and Minorities in Clinical Research
    markdown_text += "# Inclusion of Women and Minorities in Clinical Research\n\n"
    markdown_text += 'Not applicable\n\n'

    # Human Subjects Education Requirement
    markdown_text += "# Human Subjects Education Requirement\n\n"
    markdown_text += 'Not applicable\n\n'

    # Vertebrate Animals
    markdown_text += "# Vertebrate Animals\n\n"
    if component.get('uses_vertebrate_animals', False):
        markdown_text += 'No change\n\n'
    else:
        markdown_text += 'Not applicable\n\n'

    # Select Agent Research
    markdown_text += "# Select Agent Research\n\n"
    markdown_text += 'Not applicable\n\n'

    # Human Embryonic Stem Cell Line(s) Used
    markdown_text += "# Human Embryonic Stem Cell Line(s) Used\n\n"
    markdown_text += 'Not applicable\n\n'

    # Publications (only in Administrative Core)
    if component_shortname == 'Administrative Core':
        markdown_text += "# Publications\n\n"

        publications = get_components('../data/publications/publications.yaml')

        if len(publications) == 0:
            markdown_text += "Not applicable\n\n"

        for publication in publications:
            # TODO: Check dates of publications and preprint/published version

            markdown_text += f"{publication['authors']}."
            markdown_text += f" **{publication['title']}**. " 

            if 'published' in publication:
                published = publication['published']
                if published['pages'] == 'in press':
                    markdown_text += f" {published['journal']} *in press*. {published['date']}\n\n"
                    #markdown_text += f"; PubMed Central PMCID: This article is in press and has not yet received a PMCID.\n\n"
                else:
                    markdown_text += f" {published['journal']} {published['volume']}:{published['pages']}, {published['year']}. Publication {published['date']}."
                    if 'pmcid' in published:
                        markdown_text += f"; PubMed Central PMCID: {published['pmcid']}\n\n"
                    elif 'nihmsid' in published:
                        markdown_text += f"; Submitted to PubMed Central NIHMSID: {published['nihmsid']}\n\n"
                    else:
                        print(f'******* WARNING: No PMCID for {publication["title"]}')
            elif 'preprint' in publication:
                preprint = publication['preprint']
                markdown_text += f"{preprint['server']} [**Preprint**]. {preprint['date']}. Available from: {preprint['url']}\n\n" 

            markdown_text += f"*{publication['summary']}*\n\n"

            if 'projects' in publication:
                markdown_text += f"**Contributing Projects:** " + ', '.join(publication['projects']) + "\n\n"
            if 'cores' in publication:
                markdown_text += f"**Contributing Cores:** " + ', '.join(publication['cores']) + "\n\n"
            
            markdown_text += "--\n\n"
    # Project Generated Resources
    if component['type'] == 'Project' or component_shortname == 'Administrative Core':
        markdown_text += "---\n\n"    
        markdown_text += "# Project Generated Resources\n\n"
        markdown_text += "See Significant Project Generated Resources above\n\n"

    # Render markdown to document
    renderer = PythonDocxRenderer()
    render_code = MarkdownWithMath(renderer=renderer)(markdown_text)
    #for index, line in enumerate(render_code.split('\n')):
    #    print(f'{index:8d}: {line}')
    exec(render_code)

    # Save the report
    import os
    output_filename = os.path.join(output_path, component_shortname + '.docx')
    document.save(output_filename)

    # Clean up
    if os.path.exists('tmp'):
        shutil.rmtree('tmp')

if __name__ == "__main__":

    # Retrieve Project and Core metadata
    components = get_components(yaml_filename='specific-aims.yaml')

    # Generate reports for all Projects and Cores
    import os
    output_path = 'outputs'
    os.makedirs(output_path, exist_ok=True)
    for component_shortname, component in components.items():
        # Skip the Administrative Core since this is a special case we are handling manually
        #if component_shortname == 'Administrative Core':
        #    continue

        print(f"Generating report for {component['name']}")
        generate_progress_report(component_shortname, component, output_path)